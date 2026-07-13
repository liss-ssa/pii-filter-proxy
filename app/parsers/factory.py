from io import BytesIO
import fitz
from docx import Document
from app.parsers.base import ParsedDocument

def parse_document(data: bytes, content_type: str) -> ParsedDocument:
    if content_type in {"text/plain", "text/markdown", "text/html"}:
        return ParsedDocument(data.decode("utf-8-sig"), content_type)
    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(BytesIO(data))
        return ParsedDocument("\n".join(p.text for p in doc.paragraphs), content_type)
    if content_type == "application/pdf":
        pdf = fitz.open(stream=data, filetype="pdf")
        return ParsedDocument("\n".join(page.get_text("text") for page in pdf), content_type)
    raise ValueError(f"Unsupported content type: {content_type}")
